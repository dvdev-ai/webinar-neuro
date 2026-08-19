#!/usr/bin/env python3
"""Справка: цели эфира, ветки ЦА, референсы.
Формат QLAN: landscape, TNR 12/14, красная строка 709 twips, after=0, line=240.
Списки в таблицах: маркеры Word, 8 pt. Каждое предложение/фраза с заглавной буквы.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
OUT_PATHS = [
    HERE / "Справка_цели_ветки_вебинара_v1.0.docx",
]

FONT_NAME = "Times New Roman"
BLACK = RGBColor(0x00, 0x00, 0x00)
BULLET_NUM_ID = 1
TABLE_SIZE = 10
BULLET_SIZE = 8


def no_yo(text: str) -> str:
    return text.replace("ё", "е").replace("Ё", "Е")


def humanize(text: str) -> str:
    t = no_yo(str(text))
    t = t.replace("→", "->")
    t = t.replace(" — ", " - ")
    t = t.replace("—", " - ")
    return t


def cap_first(text: str) -> str:
    t = text.lstrip()
    if not t:
        return text
    return text[: len(text) - len(t)] + t[0].upper() + t[1:]


def cap_ru(text: str) -> str:
    """Заглавная в начале каждой фразы, разделенной точкой с запятой."""
    text = humanize(text).strip()
    if not text:
        return text
    if ";" not in text:
        return cap_first(text)
    parts = [cap_first(p.strip()) for p in text.split(";") if p.strip()]
    return "; ".join(parts)


def _set_run_font(run, size_pt: int, bold: bool = False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = BLACK
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_spacing_single(para):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    if spacing.get(qn("w:before")) is not None:
        del spacing.attrib[qn("w:before")]


def _set_first_line(para, twips: int | None):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if twips is None:
        if ind is not None:
            pPr.remove(ind)
        return
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLine"), str(twips))
    for attr in ("w:left", "w:right", "w:hanging"):
        key = qn(attr)
        if ind.get(key) is not None:
            del ind.attrib[key]


def ensure_bullet_numbering(doc: Document) -> None:
    try:
        numbering_part = doc.part.numbering_part
    except NotImplementedError:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.parts.numbering import NumberingPart

        numbering_part = NumberingPart.new()
        doc.part.relate_to(numbering_part, RT.NUMBERING)
    numbering = numbering_part.element
    if numbering.findall(qn("w:abstractNum")):
        return
    numbering.append(
        parse_xml(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'w:abstractNumId="0">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0">'
            '<w:start w:val="1"/>'
            '<w:numFmt w:val="bullet"/>'
            '<w:lvlText w:val="\u2022"/>'
            '<w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
            "</w:lvl>"
            "</w:abstractNum>"
        )
    )
    numbering.append(
        parse_xml(
            '<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="1">'
            '<w:abstractNumId w:val="0"/>'
            "</w:num>"
        )
    )


def _set_bullet(para, num_id: int = BULLET_NUM_ID) -> None:
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def setup_landscape(doc: Document) -> None:
    for sec in doc.sections:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)


def add_para(doc, text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, red_line=True):
    p = doc.add_paragraph()
    p.alignment = align
    _set_spacing_single(p)
    _set_first_line(p, 709 if red_line else None)
    run = p.add_run(cap_ru(text))
    _set_run_font(run, size, bold=bold)
    return p


def add_title(doc, text):
    return add_para(doc, text, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, red_line=False)


def add_h(doc, text):
    return add_para(doc, text, bold=True, size=12, red_line=True)


def add_p(doc, text):
    return add_para(doc, text, bold=False, size=12, red_line=True)


def add_li(doc, text):
    return add_para(doc, text, bold=False, size=12, red_line=True)


def _cell_clear(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    if cell.paragraphs:
        cell.paragraphs[0].clear()


def _cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=TABLE_SIZE, capitalize=True):
    _cell_clear(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    _set_spacing_single(p)
    _set_first_line(p, None)
    out = humanize(text) if not capitalize else cap_ru(text)
    run = p.add_run(out)
    _set_run_font(run, size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _cell_bullets(cell, lines, size=BULLET_SIZE):
    _cell_clear(cell)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing_single(p)
        _set_first_line(p, None)
        _set_bullet(p)
        run = p.add_run(cap_ru(line))
        _set_run_font(run, size, bold=False)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _cell_value(cell, val, *, font_size=TABLE_SIZE, bullet_size=BULLET_SIZE, capitalize=True):
    if isinstance(val, list):
        _cell_bullets(cell, val, size=bullet_size)
    else:
        _cell_text(cell, val, size=font_size, capitalize=capitalize)


def _set_table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def make_table(doc, headers, rows, col_widths=None, font_size=TABLE_SIZE, bullet_size=BULLET_SIZE, no_cap_cols=None):
    no_cap_cols = no_cap_cols or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_full_width(table)
    for cell, h in zip(table.rows[0].cells, headers):
        _cell_text(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
    for ri, row_data in enumerate(rows):
        for ci, (cell, val) in enumerate(zip(table.rows[ri + 1].cells, row_data)):
            _cell_value(cell, val, font_size=font_size, bullet_size=bullet_size, capitalize=ci not in no_cap_cols)
    if col_widths:
        _set_col_widths(table, col_widths)
    return table


def build() -> Document:
    doc = Document()
    setup_landscape(doc)
    ensure_bullet_numbering(doc)

    add_title(doc, "СПРАВКА")
    add_title(doc, "цели эфира, ветки целевой аудитории и референсы")
    add_title(doc, "вебинар по нейросетям, 90 минут")

    add_h(doc, "1. НАЗНАЧЕНИЕ ДОКУМЕНТА")
    add_p(
        doc,
        "Документ закрывает пробел в программе: зачем проводится мероприятие, что получит "
        "слушатель, что получим мы. План выступления строится от результата, а не от списка тем.",
    )

    add_h(doc, "2. ЧТО БЕРУТ ИНОСТРАННЫЕ ЭФИРЫ")
    add_p(
        doc,
        "У OpenAI, Anthropic, Stanford, Google, NIH на обложке всегда три блока: для кого; "
        "чему научит; с чем уйдет. Один эфир - одна главная целевая аудитория. "
        "Если зал смешанный, общий каркас один, а показ и следующий шаг ветвятся.",
    )

    make_table(
        doc,
        ["Эфир", "Сайт", "Чему научит", "С чем уйдет", "Для кого", "Примечание"],
        [
            [
                "OpenAI Academy. ChatGPT Foundations (60 мин)",
                "academy.openai.com",
                [
                    "ИИ и языковые модели;",
                    "Устройство ChatGPT;",
                    "Работа с запросами;",
                    "Качество результата;",
                    "Живые показы.",
                ],
                "Уверенность начать в работе и материалы",
                "Новички без технической подготовки",
                "Короткий эфир, пошаговый разбор, материалы рядом",
            ],
            [
                "NIH + OpenAI. ChatGPT 101 (60-90 мин)",
                "bioinformatics.ccr.cancer.gov/btep/classes/chatgpt-101",
                [
                    "Файлы, поиск, Canvas;",
                    "Запросы под рабочие задачи;",
                    "Безопасное использование.",
                ],
                "Навык применять инструмент в работе",
                "Сотрудники NIH",
                "Узкая аудитория, кейсы и требования под заказчика",
            ],
            [
                "Stanford. Intro to Claude",
                "uit.stanford.edu/service/techtraining/class/intro-claude",
                [
                    "Интерфейс и документы;",
                    "Углубленный поиск;",
                    "Ясные запросы и доработка;",
                    "Письмо, планирование, исследования.",
                ],
                "Уверенность в ежедневной работе",
                "Сотрудники университета",
                "Навык = действие, не теория",
            ],
            [
                "Anthropic. Agent Fundamentals",
                "anthropic.com/webinars/building-with-claude-in-europe-agent-fundamentals",
                [
                    "Сценарий и агент;",
                    "Основы на живых примерах;",
                    "Оценка результата;",
                    "Запуск в рабочей среде.",
                ],
                "Понимание, как строить агента",
                "Разработчики и IT-специалисты",
                "Образец платного IT-потока, не ядро бесплатного обзора",
            ],
            [
                "Google. Gemini for Business Analysis",
                "linkedin.com/learning/learning-google-gemini-for-business-analysis",
                [
                    "Данные и решения;",
                    "Риски и сценарии;",
                    "Этика и ограничения рядом с пользой.",
                ],
                "Навык анализа для бизнеса",
                "Бизнес-аналитики",
                "Связка роль -> задача",
            ],
            [
                "Anthropic. Claude Code Foundations",
                "anthropic.com/webinars/claude-code-foundations",
                [
                    "Установка;",
                    "Первая задача;",
                    "Правила проекта;",
                    "Масштабирование.",
                ],
                "Старт работы с кодом через Claude Code",
                "IT-специалисты",
                "Образец IT-ветки для отдельного эфира",
            ],
        ],
        col_widths=[3.4, 3.6, 4.8, 3.4, 3.2, 4.6],
        font_size=9,
        bullet_size=8,
        no_cap_cols={1},
    )

    add_h(doc, "3. ДВЕ ЦЕЛИ. ИХ НЕЛЬЗЯ СМЕШИВАТЬ")
    add_p(doc, "Цель для нас (внутренний контур). Один бесплатный эфир должен дать:")
    add_li(doc, "узнаваемость и живой контент для соцсетей;")
    add_li(doc, "базу участников с ролью при регистрации;")
    add_li(doc, "заявки на пилот (документы / парсинг) от собственников и юристов;")
    add_li(doc, "анкеты кандидатов от студентов;")
    add_li(doc, "материал для платного отраслевого потока.")
    add_p(
        doc,
        "Показатели первого цикла: регистрации, явка, вопросы, заявки в учете (пилот / платный / найм), "
        "оценка удовлетворенности.",
    )

    add_p(doc, "Цель для слушателя (это предложение эфира).")
    add_p(
        doc,
        "За 90 минут человек уходит не с пониманием темы, а с навыком и готовым результатом. "
        "Формула, как у OpenAI: уйдете с уверенностью сделать первую рабочую задачу сегодня вечером, "
        "с шаблоном запроса и планом на 7 дней.",
    )

    add_h(doc, "4. ЧТО ПОЛУЧИТ ЛЮБОЙ СЛУШАТЕЛЬ (ЯДРО)")
    add_p(doc, "К концу эфира участник сможет:")
    add_li(doc, "объяснить, чем нейросеть отличается от поиска и от обычной программы;")
    add_li(doc, "собрать рабочее ТЗ из шести частей: роль, задача, вход, формат, ограничения, критерий готовности;")
    add_li(doc, "на живом примере отличить слабый запрос от рабочего и получить черновик;")
    add_li(doc, "выбрать 1-2 инструмента под задачу, а не 12 подписок;")
    add_li(doc, "понять, какие данные нельзя отправлять в открытый чат;")
    add_li(doc, "унести: 3 шаблона запросов, чек-лист выбора инструмента, план на 7 дней.")
    add_p(doc, "Это ядро анонса. Без него целевая аудитория не собирается.")

    add_h(doc, "5. ВЕТКИ ЦЕЛЕВОЙ АУДИТОРИИ")
    add_p(doc, "Регистрация обязана спрашивать роль. Иначе ветки не работают.")

    make_table(
        doc,
        ["Ветка", "Зачем пришел", "К концу сможет", "Унесет", "Наш шаг"],
        [
            [
                "A. Собственник / руководитель МСБ",
                "Сократить рутину без лишнего найма",
                [
                    "Выбрать один процесс;",
                    "Посчитать окупаемость подписки;",
                    "Решить: бесплатно или пилот.",
                ],
                "Шаблон письма или отчета, формула цена и часы",
                "Пилот одного процесса",
            ],
            [
                "B. Юрист / клиника",
                "Быстрее входить в объем документов",
                [
                    "Собрать карту фактов по 3-5 файлам;",
                    "Требовать ссылку на файл;",
                    "Понять, где человек обязан проверить.",
                ],
                "Шаблон карты фактов, правило данных",
                "Пилот парсинга",
            ],
            [
                "C. Студент / начинающий IT-специалист",
                "Навык для портфолио и собеседования",
                [
                    "Показать цикл запрос -> черновик -> проверка;",
                    "Понять, чем Cursor отличается от чата;",
                    "Понять формат тестового задания.",
                ],
                "Шаблон запроса и мини-задача",
                "Анкета / найм",
            ],
            [
                "D. Продажи / соцсети",
                "Скорость черновиков и контента",
                [
                    "Собрать черновик поста или КП из заметок;",
                    "Держать один шаблон под канал.",
                ],
                "Шаблон контента",
                "Серия эфиров",
            ],
            [
                "E. IT / студии",
                "Агенты и код",
                "Короткий трек Cursor в показе 2",
                "Приглашение на отдельный эфир",
                "IT-поток",
            ],
        ],
        col_widths=[3.2, 4.2, 5.5, 4.0, 3.1],
        font_size=9,
        bullet_size=8,
    )

    add_h(doc, "6. ТЕКСТ АНОНСА")
    add_p(doc, "Заголовок: «Нейросети без шума: рабочий навык за 90 минут».")
    add_p(
        doc,
        "Подзаголовок: «Уйдете с шаблоном запроса, живым показом под вашу роль и планом на 7 дней. "
        "Без лекции про будущее.»",
    )
    add_p(doc, "Блок «чему научитесь»:")
    add_li(doc, "чем модель отличается от поиска и почему ответ нужно проверять;")
    add_li(doc, "как писать задание, после которого меньше правок;")
    add_li(doc, "какие инструменты брать в 2026 под текст, документы и код;")
    add_li(doc, "как не отдать данные клиента в открытый чат.")
    add_p(doc, "Блок «для кого»:")
    add_li(doc, "собственник и руководитель: время и деньги;")
    add_li(doc, "юрист и клиника: документы;")
    add_li(doc, "студент: навык и вход в команду;")
    add_li(doc, "продажи: скорость черновика.")
    add_p(doc, "Блок «с чем уйдете»:")
    add_li(doc, "3 шаблона запросов;")
    add_li(doc, "чек-лист выбора инструмента;")
    add_li(doc, "план на неделю: 1 процесс, 10 прогонов, замер времени.")

    add_h(doc, "7. ПЛАН 90 МИНУТ ОТ РЕЗУЛЬТАТА СЛУШАТЕЛЯ")
    add_p(
        doc,
        "Логика как у OpenAI Foundations и NIH 101: сначала обещание навыка, потом обучение, "
        "потом показ, который этот навык доказывает.",
    )
    make_table(
        doc,
        ["Время", "Блок", "Зачем слушателю", "Что уносит"],
        [
            ["0:00-0:06", "Рамка", "Понять, зачем здесь и что унесет", "Навык и готовый результат"],
            ["0:06-0:16", "База", "Объяснить модель и ее пределы", "Генератор, не база знаний"],
            ["0:16-0:26", "Практика 2026", "Увидеть, где уже работает", "Карта применений"],
            ["0:26-0:38", "Доступ и цена", "Выбрать контур, не слить данные", "5 критериев и правило данных"],
            ["0:38-0:50", "Инструменты", "Выбрать 1-2", "Текст, код и запасной"],
            ["0:50-0:58", "Эффект", "Замерить пользу", "План на 7 дней"],
            ["0:58-1:10", "Показ 1", "Слабый и рабочий запрос", "Шаблон ТЗ"],
            ["1:10-1:20", "Показ 2", "Свой сценарий", "Документы, код или сравнение"],
            ["1:20-1:24", "Кейс", "Понять, как выглядит пилот", "Мост к услуге"],
            ["1:24-1:28", "Вопросы", "Закрыть барьер", "3-4 ответа"],
            ["1:28-1:30", "Финал", "Следующий шаг под роль", "Материалы, пилот или найм"],
        ],
        col_widths=[2.8, 3.0, 8.5, 5.7],
        font_size=9,
        bullet_size=8,
    )
    add_p(
        doc,
        "Показ 2 не общий. Ведущий выбирает ветку по чату на старте: много собственников и юристов - "
        "документы; много IT и студентов - Cursor; зал уже все пробовал - сравнение моделей.",
    )

    add_h(doc, "8. ПЕРВЫЕ 60 СЕКУНД")
    add_p(
        doc,
        "«Сегодня не лекция про будущее. За 90 минут вы унесете три вещи: как отличить рабочий запрос "
        "от пустого, какой инструмент взять под вашу задачу, и шаблон, который можно прогнать на "
        "своей работе уже сегодня вечером. Если вы собственник - смотрите на время. "
        "Если юрист - на документы. Если студент - на навык, который можно показать.»",
    )

    add_h(doc, "9. КАРКАС СЛАЙДА 1 (КАК У OPENAI)")
    add_li(doc, "для кого (5 пунктов по ролям);")
    add_li(doc, "чему научитесь (5 пунктов, глаголы);")
    add_li(doc, "с чем уйдете (шаблоны, чек-лист, план);")
    add_li(doc, "этика и ограничения рядом с пользой, не в конце;")
    add_li(doc, "Agent Fundamentals и Claude Code - предложение платного IT-потока, не бесплатного обзора.")

    return doc


def main() -> None:
    doc = build()
    for path in OUT_PATHS:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        print("saved", path)

    d = Document(str(OUT_PATHS[0]))
    text = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    sample = next(p for p in d.paragraphs if p.text.strip().startswith("1. "))
    pPr = sample._p.find(qn("w:pPr"))
    spacing = pPr.find(qn("w:spacing"))
    ind = pPr.find(qn("w:ind"))
    bullets = 0
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr2 = p._p.find(qn("w:pPr"))
                    if pPr2 is not None and pPr2.find(qn("w:numPr")) is not None:
                        bullets += 1
    print("paras", len(d.paragraphs), "tables", len(d.tables), "bullets", bullets)
    print("has_yo", ("ё" in text) or ("Ё" in text))
    print("emdash", "\u2014" in text, "arrow", "\u2192" in text)
    print("spacing", dict(spacing.attrib) if spacing is not None else None)
    print("ind", dict(ind.attrib) if ind is not None else None)


if __name__ == "__main__":
    main()
